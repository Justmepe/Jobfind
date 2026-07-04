#!/usr/bin/env python3
"""
apply.py — generate a tailored cover-letter .docx for a job, ready to review & send.

This is the "draft for review" helper: it does NOT send anything to employers.
It writes a customized cover letter to the applications/ folder (and can email
the package to YOU for review when SMTP is configured in profile.json).

Tailoring is template-based (no API): it detects the job's flavor from the
title and selects the most relevant achievements from profile.json.

Usage:
    python apply.py --title "Senior Systems Administrator, EHS Systems" \
                    --company "Acme" --url "https://..."
    python apply.py --from-queue          # generate for every job the pipeline
                                          #   queued in applications_queue.json
    python apply.py --from-queue --email  # also email the packages to yourself
"""

import argparse
import json
import os
import re
import smtplib
import sys
from datetime import date
from email.message import EmailMessage

import requests
from docx import Document
from docx.shared import Pt

HERE = os.path.dirname(os.path.abspath(__file__))
PROFILE_FILE = os.path.join(HERE, "profile.json")
QUEUE_FILE = os.path.join(HERE, "applications_queue.json")
OUT_DIR = os.path.join(HERE, "applications")

# Map a detected job flavor to which achievement buckets to feature, in order.
FLAVOR_BUCKETS = {
    "ai": ["ai", "data", "powerplatform"],
    "systems": ["systems", "powerplatform", "data"],
    "data": ["data", "powerplatform", "ehs"],
    "powerplatform": ["powerplatform", "data", "systems"],
    "ehs": ["ehs", "powerplatform", "data"],
}

FLAVOR_PATTERNS = [
    ("ai", r"\b(ai|a\.i|ml|machine learning|llm|prompt|generative|annotat|rlhf|model)\b"),
    ("systems", r"\b(systems? admin|administrator|sharepoint|power platform|implementation)\b"),
    ("data", r"\b(data analyst|business intelligence|power bi|analytics|reporting)\b"),
    ("powerplatform", r"\b(power apps|power automate|low.?code|automation)\b"),
    ("ehs", r"\b(ehs|hse|safety|environmental|occupational|compliance|regulatory)\b"),
]

FLAVOR_FRAMING = {
    "ai": "the analytical rigor of a chemist with hands-on experience building the data pipelines and structured schemas that AI systems depend on",
    "systems": "deep, hands-on ownership of enterprise low-code systems (SharePoint, Power Apps, Power Automate) in a regulated, multi-facility environment",
    "data": "the ability to turn messy operational data into clean, auditable business-intelligence assets that leaders actually use",
    "powerplatform": "practical, production-tested Power Platform engineering delivered against real regulatory and operational requirements",
    "ehs": "a working command of OSHA, PSM, and EPA frameworks combined with the technical skill to digitize and automate them",
}


def load_profile():
    with open(PROFILE_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def detect_flavor(title):
    t = (title or "").lower()
    for flavor, pat in FLAVOR_PATTERNS:
        if re.search(pat, t):
            return flavor
    return "ehs"


def _slug(s):
    return re.sub(r"[^a-zA-Z0-9]+", "-", (s or "").strip()).strip("-")[:60] or "role"


def letter_parts(profile, title, company):
    """Return the tailored letter content (greeting, intro, bullets, close).

    Bucket selection and framing are profile-driven: a profile may supply its
    own `flavor_buckets` and `framing` maps (keyed by detected flavor). Falls
    back to the built-in defaults, then to whatever achievements exist, so the
    letter always has bullets regardless of how a profile is organized.
    """
    flavor = detect_flavor(title)
    fbuckets = profile.get("flavor_buckets") or FLAVOR_BUCKETS
    framing_map = profile.get("framing") or FLAVOR_FRAMING

    buckets = fbuckets.get(flavor) or fbuckets.get("ehs") or []
    picks = []
    for b in buckets:
        items = profile["achievements"].get(b, [])
        if items and items[0] not in picks:
            picks.append(items[0])
        if len(picks) >= 3:
            break
    # Fallback: fill from any available buckets so we never send an empty letter.
    if len(picks) < 2:
        for items in profile["achievements"].values():
            if items and items[0] not in picks:
                picks.append(items[0])
            if len(picks) >= 3:
                break

    framing = framing_map.get(flavor) or framing_map.get("ehs") or "relevant, hands-on experience"
    company = company or "your team"
    greeting = f"Dear {company} Hiring Team,"
    intro = (
        f"I am writing to apply for the {title} position at {company}. "
        f"As {profile['current_role']}, I am {profile['one_liner']}. "
        f"What I would bring to this role is {framing}."
    )
    lead = "A few examples of the work I have contributed to that map directly to this role:"
    bullets = [p[0].upper() + p[1:] + "." for p in picks]
    close = (
        f"I am excited by the opportunity at {company} and would welcome the chance to "
        f"discuss how my mix of regulatory knowledge, systems-building, and data skills "
        f"can contribute. Thank you for your time and consideration."
    )
    return {
        "flavor": flavor,
        "greeting": greeting,
        "intro": intro,
        "lead": lead,
        "bullets": bullets,
        "close": close,
    }


def build_cover_letter(profile, title, company):
    parts = letter_parts(profile, title, company)

    doc = Document()
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(11)

    h = doc.add_paragraph()
    r = h.add_run(profile["name"])
    r.bold = True
    r.font.size = Pt(15)
    contact_bits = [profile.get("location"), profile.get("email"),
                    profile.get("phone"), profile.get("linkedin")]
    contact = "  |  ".join([c for c in contact_bits if c])
    doc.add_paragraph(contact)
    doc.add_paragraph(date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()
    doc.add_paragraph(parts["greeting"])
    doc.add_paragraph(parts["intro"])
    doc.add_paragraph(parts["lead"])
    for line in parts["bullets"]:
        doc.add_paragraph(style="List Bullet").add_run(line)
    doc.add_paragraph(parts["close"])
    doc.add_paragraph()
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph(profile["name"])

    os.makedirs(OUT_DIR, exist_ok=True)
    fname = f"CoverLetter_{_slug(company)}_{_slug(title)}.docx"
    path = os.path.join(OUT_DIR, fname)
    doc.save(path)
    return path, parts["flavor"]


def build_application_message(profile, title, company, url="", to_email=""):
    """Build the full application email (body + cover letter + resume attached).

    Returns (EmailMessage, has_resume). `to_email` may be empty for a draft.
    """
    parts = letter_parts(profile, title, company)
    cover_path, _ = build_cover_letter(profile, title, company)
    resume = _find_resume(profile)

    contact_bits = [profile.get("location"), profile.get("email"),
                    profile.get("phone"), profile.get("linkedin")]
    contact = " | ".join([c for c in contact_bits if c])

    plain = [parts["greeting"], "", parts["intro"], "", parts["lead"], ""]
    plain += [f"  • {b}" for b in parts["bullets"]]
    plain += ["", parts["close"], "", "Sincerely,", profile["name"], contact]
    if url:
        plain += ["", f"(In reference to your posting: {url})"]

    bullets_html = "".join(f'<li style="margin:6px 0">{b}</li>' for b in parts["bullets"])
    posting = (
        f'<p style="color:#8b949e;font-size:12px">In reference to your posting: '
        f'<a href="{url}">{url}</a></p>' if url else ""
    )
    html = f"""<!DOCTYPE html><html><body style="margin:0;background:#f6f8fa">
      <div style="font-family:Georgia,'Times New Roman',serif;max-width:640px;margin:0 auto;
                  padding:28px;background:#fff;color:#222;line-height:1.6;font-size:15px">
        <p>{parts['greeting']}</p>
        <p>{parts['intro']}</p>
        <p>{parts['lead']}</p>
        <ul style="padding-left:20px">{bullets_html}</ul>
        <p>{parts['close']}</p>
        <p style="margin-bottom:2px">Sincerely,</p>
        <p style="margin-top:2px"><b>{profile['name']}</b><br>
           <span style="color:#555;font-size:13px">{contact}</span></p>
        {posting}
      </div>
    </body></html>"""

    msg = EmailMessage()
    msg["Subject"] = f"Application for {title} — {profile['name']}"
    msg["From"] = f"{profile['name']} <{profile['email']}>"
    if to_email:
        msg["To"] = to_email
    msg["Reply-To"] = profile["email"]
    msg.set_content("\n".join(plain))
    msg.add_alternative(html, subtype="html")

    for path in [cover_path, resume]:
        if path and os.path.exists(path):
            _attach_file(msg, path)
    return msg, bool(resume)


def send_application(profile, title, company, to_email, url=""):
    """Send an application email TO an employer. OUTWARD-FACING — use deliberately."""
    smtp = profile.get("smtp") or {}
    if not smtp.get("app_password"):
        sys.exit("Cannot send: no smtp.app_password in profile.json")
    msg, has_resume = build_application_message(profile, title, company, url, to_email)
    with smtplib.SMTP_SSL(smtp.get("host", "smtp.gmail.com"), int(smtp.get("port", 465))) as s:
        s.login(profile["email"], smtp["app_password"])
        s.send_message(msg)
    print(f"  APPLICATION SENT to {to_email} "
          f"(cover letter{' + resume' if has_resume else ''} attached)")


def create_gmail_drafts(profile, jobs):
    """Create a ready-to-send Gmail draft per job in the Drafts folder (via IMAP).

    Each draft is a complete application (body + cover letter + resume). The
    recipient is pre-filled only if the job carries an application email;
    otherwise you add it (or apply via the posting's portal) before sending.
    """
    smtp = profile.get("smtp") or {}
    if not smtp.get("app_password"):
        sys.exit("Cannot create drafts: no smtp.app_password in profile.json")
    import imaplib
    import time

    M = imaplib.IMAP4_SSL("imap.gmail.com", 993)
    M.login(profile["email"], smtp["app_password"])
    made = 0
    for j in jobs:
        to_email = j.get("apply_email", "")  # usually empty (portal jobs)
        msg, _ = build_application_message(
            profile, j["title"], j.get("company", ""), j.get("url", ""), to_email
        )
        M.append(
            '"[Gmail]/Drafts"', "\\Draft",
            imaplib.Time2Internaldate(time.time()),
            msg.as_bytes(),
        )
        made += 1
        print(f"  draft created: {j['title']} @ {j.get('company','')}"
              + (f"  -> {to_email}" if to_email else "  (add recipient)"))
    M.logout()
    print(f"  {made} Gmail draft(s) created in your Drafts folder")
    return made


DOCX_MIME = "vnd.openxmlformats-officedocument.wordprocessingml.document"


def _attach_file(msg, path):
    """Attach a file, choosing the MIME subtype from its extension."""
    ext = os.path.splitext(path)[1].lower()
    subtype = "pdf" if ext == ".pdf" else DOCX_MIME
    with open(path, "rb") as f:
        msg.add_attachment(
            f.read(), maintype="application", subtype=subtype,
            filename=os.path.basename(path),
        )


def _find_resume(profile):
    """Locate the resume file — check the app folder first, then its parent."""
    name = profile.get("resume_file", "")
    if not name:
        return None
    for cand in (os.path.join(HERE, name), os.path.join(os.path.dirname(HERE), name)):
        if os.path.exists(cand):
            return cand
    return None


def _email_html(jobs_done, resume_attached):
    cards = ""
    for j in jobs_done:
        url = j.get("url", "")
        link = (
            f'<a href="{url}" style="color:#1f6feb;text-decoration:none">View posting &rarr;</a>'
            if url
            else ""
        )
        src = j.get("source", "")
        meta = " &middot; ".join([x for x in [j.get("company", ""), src] if x])
        cards += f"""
        <div style="border:1px solid #e1e4e8;border-radius:8px;padding:14px 16px;margin:12px 0;background:#fff">
          <div style="font-weight:700;font-size:15px;color:#111">{j['title']}</div>
          <div style="color:#555;font-size:13px;margin:4px 0 8px">{meta}</div>
          {link}
        </div>"""
    resume_note = (
        "Your resume is attached alongside each cover letter."
        if resume_attached
        else "<b>Note:</b> your resume was not found on the server, so only the cover letter is attached."
    )
    return f"""<!DOCTYPE html>
<html><body style="margin:0;padding:0;background:#f6f8fa">
  <div style="font-family:-apple-system,Segoe UI,Arial,sans-serif;max-width:640px;margin:0 auto;padding:24px">
    <h2 style="color:#1f6feb;margin:0 0 4px">📄 {len(jobs_done)} application draft{'s' if len(jobs_done)!=1 else ''} ready to review</h2>
    <p style="color:#444;font-size:14px;line-height:1.5;margin:8px 0 4px">
      Tailored cover letters are attached as Word documents. {resume_note}
      Review each one, then send it to the employer.
    </p>
    {cards}
    <p style="color:#8b949e;font-size:12px;margin-top:20px;border-top:1px solid #e1e4e8;padding-top:12px">
      Generated automatically by jobwatch. These are drafts — always review before sending.
    </p>
  </div>
</body></html>"""


def email_package(profile, jobs_done):
    """Email the generated packages to yourself for review (optional)."""
    smtp = profile.get("smtp") or {}
    if not smtp.get("app_password"):
        print("  (skip email: no smtp.app_password in profile.json)")
        return
    msg = EmailMessage()
    n = len(jobs_done)
    msg["Subject"] = f"[jobwatch] {n} application draft{'s' if n != 1 else ''} ready to review"
    msg["From"] = f"jobwatch <{profile['email']}>"
    msg["To"] = profile["email"]

    resume = _find_resume(profile)

    # Plain-text fallback, then the HTML alternative.
    plain = ["Your tailored cover-letter drafts are attached. Review, then send:", ""]
    for j in jobs_done:
        plain.append(f"- {j['title']} @ {j.get('company','')}")
        if j.get("url"):
            plain.append(f"  {j['url']}")
    msg.set_content("\n".join(plain))
    msg.add_alternative(_email_html(jobs_done, bool(resume)), subtype="html")

    # Attachments go on the top-level (multipart/mixed) message.
    for j in jobs_done:
        _attach_file(msg, j["path"])
    if resume:
        _attach_file(msg, resume)

    host = smtp.get("host", "smtp.gmail.com")
    port = int(smtp.get("port", 465))
    with smtplib.SMTP_SSL(host, port) as s:
        s.login(profile["email"], smtp["app_password"])
        s.send_message(msg)
    print(
        f"  emailed {n} draft(s) to {profile['email']} "
        f"({'resume attached' if resume else 'RESUME MISSING'})"
    )


def notify_discord(profile, jobs_done, mode="email"):
    """Post a heads-up to Discord that drafts are waiting."""
    cfg_file = os.path.join(HERE, "config.json")
    webhook = os.environ.get("DISCORD_WEBHOOK_URL")
    if not webhook and os.path.exists(cfg_file):
        try:
            webhook = json.load(open(cfg_file, encoding="utf-8")).get("discord_webhook_url")
        except Exception:
            webhook = None
    if not webhook:
        return
    intro = (
        "Ready in your **Gmail Drafts** — open each, add the recipient if the "
        "posting lists one, and send."
        if mode == "drafts"
        else f"Emailed to **{profile['email']}** — review each and send."
    )
    fields = []
    for j in jobs_done[:10]:
        company = (j.get("company") or "—").strip()
        src = (j.get("source") or "").split(" (")[0]
        value = f"🏢 {company}" + (f"  ·  {src}" if src else "")
        if j.get("url"):
            value += f"\n🔗 [View posting]({j['url']})"
        fields.append({
            "name": j["title"][:250],
            "value": value[:1024],
            "inline": False,
        })
    n = len(jobs_done)
    embed = {
        "title": f"📄 {n} application draft{'s' if n != 1 else ''} ready",
        "description": intro,
        "color": 15844367,
        "fields": fields,
        "footer": {"text": "Review before sending"
                   + (f"  ·  +{n - 10} more" if n > 10 else "")},
    }
    try:
        requests.post(webhook, json={"embeds": [embed]}, timeout=20)
    except Exception as e:
        print(f"  (discord notify failed: {e})")


def main():
    ap = argparse.ArgumentParser(description="Generate tailored cover-letter drafts.")
    ap.add_argument("--title")
    ap.add_argument("--company")
    ap.add_argument("--url", default="")
    ap.add_argument("--from-queue", action="store_true",
                    help="Generate for every job in applications_queue.json")
    ap.add_argument("--email", action="store_true",
                    help="Also email the packages to yourself (needs smtp in profile.json)")
    ap.add_argument("--send-to", metavar="EMAIL",
                    help="OUTWARD-FACING: send the application directly to this "
                         "employer address (with --title/--company).")
    ap.add_argument("--drafts", action="store_true",
                    help="Create a ready-to-send Gmail draft per job (in your "
                         "Drafts folder) instead of emailing drafts to yourself.")
    args = ap.parse_args()

    profile = load_profile()

    # Outward-facing application send to an employer address.
    if args.send_to:
        if not args.title:
            sys.exit("--send-to requires --title (and usually --company).")
        send_application(profile, args.title, args.company or "",
                         args.send_to, args.url)
        return

    if args.from_queue:
        if not os.path.exists(QUEUE_FILE):
            sys.exit("No applications_queue.json yet — run the pipeline first.")
        with open(QUEUE_FILE, "r", encoding="utf-8") as f:
            jobs = json.load(f)
    elif args.title:
        jobs = [{"title": args.title, "company": args.company or "", "url": args.url}]
    else:
        sys.exit("Provide --title/--company or --from-queue.")

    # Gmail-draft mode: one ready-to-send draft per job in your Drafts folder.
    if args.drafts:
        if jobs:
            create_gmail_drafts(profile, jobs)
            notify_discord(profile, jobs, mode="drafts")
            if args.from_queue:
                with open(QUEUE_FILE, "w", encoding="utf-8") as f:
                    json.dump([], f)
                print("  queue cleared")
        else:
            print("No jobs to draft.")
        return

    done = []
    for j in jobs:
        path, flavor = build_cover_letter(profile, j["title"], j.get("company"))
        print(f"[{flavor:12}] {os.path.basename(path)}")
        done.append({**j, "path": path})

    if args.email and done:
        email_package(profile, done)
        notify_discord(profile, done)
        # Clear the queue after a successful email run so drafts aren't re-sent.
        if args.from_queue and (profile.get("smtp") or {}).get("app_password"):
            with open(QUEUE_FILE, "w", encoding="utf-8") as f:
                json.dump([], f)
            print("  queue cleared")

    print(f"\nGenerated {len(done)} cover letter(s) in {OUT_DIR}")


if __name__ == "__main__":
    main()
